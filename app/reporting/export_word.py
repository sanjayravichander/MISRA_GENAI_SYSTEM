#!/usr/bin/env python3
"""
export_word.py  —  Milestone 2: Word report exporter (pure Python, no Node.js).
Generates a professional .docx from evaluated_fixes.json + audit.json.

Usage:
    python export_word.py <evaluated_fixes.json> <audit.json> <output.docx>

Example:
    python export_word.py data\\output\\run_001\\evaluated_fixes.json ^
                          data\\audit\\run_001_audit.json             ^
                          data\\output\\run_001\\misra_report.docx
"""

from __future__ import annotations

import json
import os
import sys
from datetime import datetime
from typing import Any, Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches


# ---------------------------------------------------------------------------
# Palette
# ---------------------------------------------------------------------------

BRAND_BLUE  = RGBColor(0x1F, 0x49, 0x7D)
BRAND_MID   = RGBColor(0x2E, 0x75, 0xB6)
GREY_HEAD   = RGBColor(0x40, 0x40, 0x40)
GREY_BODY   = RGBColor(0x22, 0x22, 0x22)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
GREEN       = RGBColor(0x1E, 0x7B, 0x34)
AMBER       = RGBColor(0xD0, 0x78, 0x00)
RED         = RGBColor(0xC0, 0x39, 0x2B)
CODE_GREY   = RGBColor(0x1A, 0x1A, 0x1A)
CODE_BG     = "F2F2F2"
HEADER_BG   = "1F497D"
ROW_ALT     = "EEF3FA"
LIGHT_BG    = "D9E2F3"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def load_json(path: str) -> dict:
    with open(path, encoding="utf-8") as f:
        return json.load(f)


def fmt_duration(seconds: float) -> str:
    m, s = divmod(int(seconds), 60)
    return f"{m}m {s}s" if m else f"{s}s"


def fmt_ts(iso: str) -> str:
    try:
        dt = datetime.fromisoformat(iso)
        return dt.strftime("%d %b %Y  %H:%M:%S")
    except Exception:
        return iso


def conf_color(conf: str) -> RGBColor:
    return {"High": GREEN, "Medium": AMBER, "Low": RED}.get(conf, GREY_HEAD)


def risk_color(risk: str) -> RGBColor:
    return {"Low": GREEN, "Medium": AMBER, "High": RED,
            "Critical": RED}.get(risk, GREY_HEAD)


def clean_code(snippet: str) -> str:
    s = snippet.strip()
    for fence in ("```c", "```cpp", "```", "~~~"):
        s = s.replace(fence, "")
    return s.strip()


def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def add_horizontal_rule(doc: Document, color: str = "2E75B6"):
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    return p


def add_header_footer(doc: Document, run_id: str):
    section = doc.sections[0]

    # Header
    header = section.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.clear()
    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "2E75B6")
    pBdr.append(bot)
    pPr.append(pBdr)
    run = hp.add_run(f"MISRA C Compliance Report  |  Run: {run_id}")
    run.font.size      = Pt(8)
    run.font.color.rgb = GREY_HEAD

    # Footer
    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    pPr2 = fp._p.get_or_add_pPr()
    pBdr2 = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"),   "single")
    top.set(qn("w:sz"),    "6")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), "2E75B6")
    pBdr2.append(top)
    pPr2.append(pBdr2)

    tabs = OxmlElement("w:tabs")
    tab  = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), "9360")
    tabs.append(tab)
    pPr2.append(tabs)

    lr = fp.add_run("MISRA GenAI System  \u2014  Confidential")
    lr.font.size      = Pt(8)
    lr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    fp.add_run("\t")

    pr = fp.add_run()
    pr.font.size      = Pt(8)
    pr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    pr._r.append(fld)
    instr = OxmlElement("w:instrText")
    instr.text = " PAGE "
    pr._r.append(instr)
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    pr._r.append(fld2)


# ---------------------------------------------------------------------------
# Cover page
# ---------------------------------------------------------------------------

def build_cover(doc: Document, run_id, started_at, completed_at,
                total_dur, excel_src):
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MISRA C Compliance Report")
    r.font.size      = Pt(28)
    r.font.bold      = True
    r.font.color.rgb = BRAND_BLUE
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Automated Static Analysis & Fix Suggestions")
    r.font.size      = Pt(14)
    r.font.color.rgb = BRAND_MID
    p.paragraph_format.space_after = Pt(12)

    add_horizontal_rule(doc)

    for label, value in [("Run ID", run_id), ("Started", started_at),
                          ("Completed", completed_at), ("Duration", total_dur)]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{label}:  {value}")
        r.font.size      = Pt(11)
        r.font.color.rgb = GREY_HEAD
        p.paragraph_format.space_after = Pt(3)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Source: {excel_src}")
    r.font.size   = Pt(9)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    for _ in range(8):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Generated by MISRA GenAI System")
    r.font.size   = Pt(9)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Tables
# ---------------------------------------------------------------------------

def add_two_col_table(doc: Document, rows: list, col_widths=(3.5, 3.0)):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    for i, row_data in enumerate(rows):
        col0, col1 = row_data[0], row_data[1]
        val_color  = row_data[2] if len(row_data) > 2 else None

        row = table.add_row()
        for ci, (cell, text) in enumerate(
                [(row.cells[0], col0), (row.cells[1], col1)]):
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(str(text))
            run.font.size = Pt(10)
            if i == 0:
                set_cell_bg(cell, HEADER_BG)
                run.font.bold      = True
                run.font.color.rgb = WHITE
            else:
                set_cell_bg(cell, ROW_ALT if i % 2 == 0 else "FFFFFF")
                if ci == 1 and val_color:
                    run.font.color.rgb = val_color
                    run.font.bold      = True
                else:
                    run.font.color.rgb = GREY_BODY


def add_four_col_table(doc: Document, rows: list,
                       col_widths=(0.7, 2.5, 1.1, 2.2)):
    table = doc.add_table(rows=0, cols=4)
    table.style = "Table Grid"
    n = len(rows)

    for i, row_data in enumerate(rows):
        row = table.add_row()
        for ci, (cell, text) in enumerate(zip(row.cells, row_data)):
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(str(text))
            run.font.size = Pt(9)
            if i == 0:
                set_cell_bg(cell, HEADER_BG)
                run.font.bold      = True
                run.font.color.rgb = WHITE
            elif i == n - 1:
                set_cell_bg(cell, LIGHT_BG)
                run.font.bold      = True
                run.font.color.rgb = GREY_HEAD
            else:
                set_cell_bg(cell, ROW_ALT if i % 2 == 0 else "FFFFFF")
                run.font.color.rgb = GREY_BODY


# ---------------------------------------------------------------------------
# Executive summary
# ---------------------------------------------------------------------------

def build_summary(doc: Document, fixes_data: dict, audit: dict):
    p = doc.add_heading("Executive Summary", level=1)
    p.runs[0].font.color.rgb = BRAND_BLUE
    add_horizontal_rule(doc)

    w_count   = fixes_data.get("warning_count", 0)
    c_high    = fixes_data.get("confidence_high", 0)
    c_medium  = fixes_data.get("confidence_medium", 0)
    c_low     = fixes_data.get("confidence_low", 0)
    manual    = fixes_data.get("needs_manual_review", 0)
    corrected = fixes_data.get("fixes_corrected", 0)

    quality_rows = [
        ["Metric", "Value"],
        ["Total Warnings",                    str(w_count),   None],
        ["High Confidence",                   str(c_high),    GREEN],
        ["Medium Confidence",                 str(c_medium),  AMBER],
        ["Low Confidence",                    str(c_low),     RED],
        ["Needs Manual Review",               str(manual),    RED if manual > 0 else GREEN],
        ["Fixes Auto-corrected by Evaluator", str(corrected), None],
    ]
    add_two_col_table(doc, quality_rows)
    doc.add_paragraph()

    p = doc.add_heading("Pipeline Timing", level=2)
    p.runs[0].font.color.rgb = BRAND_MID

    phases = audit.get("phases", {})
    p6a = phases.get("6a", {})
    p6b = phases.get("6b", {})
    p7  = phases.get("7",  {})
    p8  = phases.get("8",  {})

    timing_rows = [
        ("Phase", "Description", "Duration", "Notes"),
        ("6a", "Parse Polyspace Excel report",
         fmt_duration(p6a.get("duration_s", 0)),
         f"{p6a.get('warnings','?')} warnings"),
        ("6b", "FAISS retrieval — MISRA context",
         fmt_duration(p6b.get("duration_s", 0)), "—"),
        ("7",  "Generate fix suggestions (LLM)",
         fmt_duration(p7.get("duration_s", 0)),
         f"{p7.get('fixes_generated','?')} generated, "
         f"{p7.get('parse_errors','?')} errors"),
        ("8",  "Evaluate & self-critique (LLM)",
         fmt_duration(p8.get("duration_s", 0)),
         f"{p8.get('fixes_corrected','?')} corrected, "
         f"{p8.get('needs_manual_review','?')} flagged"),
        ("Total", "End-to-end pipeline",
         fmt_duration(audit.get("total_duration_s", 0)), "—"),
    ]
    add_four_col_table(doc, timing_rows)
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Warning section
# ---------------------------------------------------------------------------

def build_warning_section(doc: Document, w: Dict[str, Any]):
    conf   = (w.get("overall_confidence") or
              (w.get("evaluation") or {}).get("overall_confidence", "—"))
    manual = (w.get("needs_manual_review") or
              (w.get("evaluation") or {}).get("needs_manual_review", False))

    p = doc.add_heading(
        f"{w.get('warning_id','')}  \u2014  {w.get('rule_id','')}", level=2)
    p.runs[0].font.color.rgb = BRAND_MID

    p = doc.add_paragraph()
    r = p.add_run(f"Overall confidence: {conf}")
    r.font.bold      = True
    r.font.size      = Pt(10)
    r.font.color.rgb = conf_color(conf)
    if manual:
        r2 = p.add_run("    \u26a0 MANUAL REVIEW REQUIRED")
        r2.font.bold      = True
        r2.font.size      = Pt(10)
        r2.font.color.rgb = RED
    p.paragraph_format.space_after = Pt(6)

    def label_para(text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.bold      = True
        r.font.size      = Pt(10)
        r.font.color.rgb = GREY_HEAD
        p.paragraph_format.space_after = Pt(2)

    def body_para(text, italic=False, indent=0):
        p = doc.add_paragraph(str(text))
        p.runs[0].font.size      = Pt(10)
        p.runs[0].font.italic    = italic
        p.runs[0].font.color.rgb = GREY_BODY
        p.paragraph_format.space_after = Pt(6)
        if indent:
            p.paragraph_format.left_indent = Inches(indent)
        return p

    if w.get("explanation"):
        label_para("Explanation")
        body_para(w["explanation"])

    if w.get("risk_analysis"):
        label_para("Risk Analysis")
        body_para(w["risk_analysis"])

    label_para("Ranked Fixes")

    for fix in w.get("ranked_fixes", []):
        rank  = fix.get("rank", "?")
        desc  = fix.get("description", "")
        code  = clean_code(fix.get("code_change", ""))
        rat   = fix.get("rationale", "")
        risk  = fix.get("risk_level", "")
        fconf = fix.get("confidence", "")
        issues= fix.get("issues_found", [])
        was_corrected = fix.get("was_corrected", False)

        p = doc.add_paragraph()
        r = p.add_run(f"Fix {rank}  ")
        r.font.bold      = True
        r.font.size      = Pt(10)
        r.font.color.rgb = BRAND_MID
        if fconf:
            r2 = p.add_run(f"Conf: {fconf}  ")
            r2.font.bold      = True
            r2.font.size      = Pt(9)
            r2.font.color.rgb = conf_color(fconf)
        if risk:
            r3 = p.add_run(f"Risk: {risk}")
            r3.font.bold      = True
            r3.font.size      = Pt(9)
            r3.font.color.rgb = risk_color(risk)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(2)

        if desc:
            p = doc.add_paragraph(desc)
            p.runs[0].font.italic    = True
            p.runs[0].font.size      = Pt(9)
            p.runs[0].font.color.rgb = GREY_HEAD
            p.paragraph_format.space_after = Pt(3)

        if code:
            p = doc.add_paragraph()
            p.add_run("Code change:").font.bold = True
            p.runs[0].font.size = Pt(9)
            p.paragraph_format.space_after = Pt(2)
            for line in code.splitlines():
                cp = doc.add_paragraph(line or " ")
                cp.runs[0].font.name      = "Courier New"
                cp.runs[0].font.size      = Pt(8)
                cp.runs[0].font.color.rgb = CODE_GREY
                cp.paragraph_format.left_indent  = Inches(0.2)
                cp.paragraph_format.space_before = Pt(0)
                cp.paragraph_format.space_after  = Pt(0)
                pPr = cp._p.get_or_add_pPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"),   "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"),  CODE_BG)
                pPr.append(shd)
            sp = doc.add_paragraph()
            sp.paragraph_format.space_after = Pt(2)

        if rat:
            p = doc.add_paragraph(f"Rationale: {rat}")
            p.runs[0].font.italic    = True
            p.runs[0].font.size      = Pt(9)
            p.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            p.paragraph_format.left_indent = Inches(0.15)
            p.paragraph_format.space_after = Pt(3)

        if issues:
            p = doc.add_paragraph()
            r = p.add_run("Issues found:")
            r.font.bold      = True
            r.font.size      = Pt(9)
            r.font.color.rgb = RED
            for iss in issues:
                ip = doc.add_paragraph(f"\u2022 {iss}")
                ip.runs[0].font.size      = Pt(9)
                ip.runs[0].font.color.rgb = RED
                ip.paragraph_format.left_indent = Inches(0.2)

        if was_corrected:
            p = doc.add_paragraph("\u2714 Evaluator corrected this fix")
            p.runs[0].font.bold      = True
            p.runs[0].font.size      = Pt(9)
            p.runs[0].font.color.rgb = GREEN
            p.paragraph_format.space_after = Pt(3)

    notes = (w.get("evaluator_notes") or
             (w.get("evaluation") or {}).get("evaluator_notes", ""))
    if notes:
        label_para("Evaluator Notes")
        body_para(notes, italic=True, indent=0.15)

    if w.get("deviation_advice"):
        label_para("Deviation Advice")
        body_para(w["deviation_advice"])

    add_horizontal_rule(doc, color="CCCCCC")


# ---------------------------------------------------------------------------
# Appendix
# ---------------------------------------------------------------------------

def build_appendix(doc: Document, results: List[Dict]):
    doc.add_page_break()
    p = doc.add_heading(
        "Appendix \u2014 Warnings Requiring Manual Review", level=1)
    p.runs[0].font.color.rgb = BRAND_BLUE
    add_horizontal_rule(doc)

    manual_warnings = [
        w for w in results
        if w.get("needs_manual_review") or
           (w.get("evaluation") or {}).get("needs_manual_review", False)
    ]

    if not manual_warnings:
        p = doc.add_paragraph("No warnings require manual review.")
        p.runs[0].font.italic    = True
        p.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    else:
        for w in manual_warnings:
            notes = (w.get("evaluator_notes") or
                     (w.get("evaluation") or {}).get("evaluator_notes", ""))
            p = doc.add_heading(
                f"{w.get('warning_id','')}  \u2014  {w.get('rule_id','')}", level=3)
            p.runs[0].font.color.rgb = GREY_HEAD
            if notes:
                p = doc.add_paragraph(notes)
                p.runs[0].font.italic    = True
                p.runs[0].font.size      = Pt(9)
                p.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            add_horizontal_rule(doc, color="CCCCCC")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    if len(sys.argv) < 4:
        print("Usage: python export_word.py <evaluated_fixes.json> "
              "<audit.json> <output.docx>")
        sys.exit(1)

    fixes_path = sys.argv[1]
    audit_path = sys.argv[2]
    out_path   = os.path.abspath(sys.argv[3])

    fixes_data   = load_json(fixes_path)
    audit        = load_json(audit_path)
    run_id       = audit.get("run_id", "—")
    started_at   = fmt_ts(audit.get("started_at", ""))
    completed_at = fmt_ts(audit.get("completed_at", ""))
    total_dur    = fmt_duration(audit.get("total_duration_s", 0))
    excel_src    = audit.get("excel_report", "—")
    results      = fixes_data.get("results", [])

    doc     = Document()
    section = doc.sections[0]
    section.page_width    = Inches(8.5)
    section.page_height   = Inches(11)
    section.left_margin   = Inches(0.75)
    section.right_margin  = Inches(0.75)
    section.top_margin    = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)

    add_header_footer(doc, run_id)
    build_cover(doc, run_id, started_at, completed_at, total_dur, excel_src)
    build_summary(doc, fixes_data, audit)

    p = doc.add_heading("Warning Details", level=1)
    p.runs[0].font.color.rgb = BRAND_BLUE
    add_horizontal_rule(doc)
    for w in results:
        build_warning_section(doc, w)

    build_appendix(doc, results)

    doc.save(out_path)
    size_kb = os.path.getsize(out_path) // 1024
    print(f"[OK] Word report written \u2192 {out_path}  ({size_kb} KB)")


if __name__ == "__main__":
    main()